#!/usr/bin/env python3
"""
Build print-ready .docx and .pdf from utkast markdown files.
Handles: BibTeX citation rendering, math subscript/superscript,
title page, page numbers, proper typography.
"""

import re
import sys
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from weasyprint import HTML

BASE_DIR = "/workspaces/NBNO.py/formgjevars traktat/redaksjon"
BIB_FILE = "/workspaces/NBNO.py/formgjevars traktat/references.bib"
INPUT = os.path.join(BASE_DIR, "utkast_02.md")
OUTPUT_DOCX = os.path.join(BASE_DIR, "utkast_02.docx")
OUTPUT_PDF = os.path.join(BASE_DIR, "utkast_02.pdf")

# ---------------------------------------------------------------------------
# BibTeX parser (minimal, for author-year citations)
# ---------------------------------------------------------------------------

def parse_bibtex(path):
    """Parse .bib file into dict of key -> {author, year, title}."""
    entries = {}
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()

    # Match each @type{key, ... }
    pattern = re.compile(r'@\w+\{([^,]+),\s*(.*?)\n\}', re.DOTALL)
    for m in pattern.finditer(text):
        key = m.group(1).strip()
        body = m.group(2)
        entry = {}
        for field_m in re.finditer(r'(\w+)\s*=\s*\{(.*?)\}', body, re.DOTALL):
            fname = field_m.group(1).lower()
            fval = field_m.group(2).strip()
            # Clean LaTeX commands
            fval = re.sub(r'\\textit\{([^}]+)\}', r'\1', fval)
            fval = re.sub(r'\{\\[a-z]+\s+([^}]*)\}', r'\1', fval)
            fval = re.sub(r'[{}]', '', fval)
            fval = re.sub(r'\\[a-z]+', '', fval)
            entry[fname] = fval.strip()
        entries[key] = entry
    return entries


def make_citation_label(entry):
    """Create (Author, Year) or (Title, Year) label."""
    author = entry.get('author', '')
    year = entry.get('year', '')
    title = entry.get('title', '')

    if author:
        # Get first author surname
        parts = author.split(' and ')
        first = parts[0].strip()
        if ',' in first:
            surname = first.split(',')[0].strip()
        else:
            surname = first.split()[-1] if first.split() else first
        if 'others' in author and len(parts) > 1:
            surname += ' et al.'
        elif len(parts) > 2:
            surname += ' et al.'
        elif len(parts) == 2:
            second = parts[1].strip()
            if ',' in second:
                s2 = second.split(',')[0].strip()
            else:
                s2 = second.split()[-1] if second.split() else second
            if s2 != 'others':
                surname += ' og ' + s2
            else:
                surname += ' et al.'
    else:
        # Use short title
        surname = title[:40] + ('...' if len(title) > 40 else '')

    if year:
        return f'({surname}, {year})'
    return f'({surname})'


def resolve_citations(text, bib):
    """Replace [bibtex_key] and [key1; key2] with (Author, Year) labels."""
    def replace_cite(m):
        keys_str = m.group(1)
        keys = [k.strip() for k in keys_str.split(';')]
        labels = []
        for key in keys:
            if key in bib:
                labels.append(make_citation_label(bib[key]))
            elif key.startswith('M:'):
                return m.group(0)  # Keep [M: ...] cross-references
            else:
                labels.append(f'[{key}]')
        return '; '.join(labels)

    # Match [something] but not [M: ...] and not [Definisjon] etc.
    text = re.sub(
        r'\[([a-z][a-z0-9_]+(?:;\s*[a-z][a-z0-9_]+)*)\]',
        replace_cite,
        text
    )
    return text


# ---------------------------------------------------------------------------
# Math notation: convert M_C -> M꜀, R^n -> Rⁿ, etc. using Unicode
# ---------------------------------------------------------------------------

# Some subscript chars don't exist in Unicode, use best alternatives
SUB_MAP = {
    'C': '\u1D04', 'c': '\u1D04',  # small cap C
    'i': '\u1D62', 'j': '\u2C7C', 'k': '\u2096',
    'm': '\u2098', 'n': '\u2099', 's': '\u209B', 't': '\u209C',
    'x': '\u2093', 'a': '\u2090', 'e': '\u2091', 'o': '\u2092',
    'r': '\u1D63', 'u': '\u1D64',
    '0': '\u2080', '1': '\u2081', '2': '\u2082', '3': '\u2083',
    '4': '\u2084', '5': '\u2085', '6': '\u2086', '7': '\u2087',
    '8': '\u2088', '9': '\u2089',
}
SUP_MAP = {
    'n': '\u207F', 'i': '\u2071', 'k': '\u1D4F',
    '0': '\u2070', '1': '\u00B9', '2': '\u00B2', '3': '\u00B3',
    '4': '\u2074', '5': '\u2075', '6': '\u2076', '7': '\u2077',
    '8': '\u2078', '9': '\u2079',
    '+': '\u207A', '-': '\u207B', '=': '\u207C',
    '(': '\u207D', ')': '\u207E',
}

def to_subscript(s):
    return ''.join(SUB_MAP.get(c, c) for c in s)

def to_superscript(s):
    return ''.join(SUP_MAP.get(c, c) for c in s)


def convert_math_notation(text):
    """Convert _X subscripts and ^X superscripts to Unicode equivalents."""
    # Handle patterns like R^n, R^k, sigma^2
    text = re.sub(r'([A-Za-z\w])\^(\{[^}]+\}|[A-Za-z0-9])', lambda m: m.group(1) + to_superscript(m.group(2).strip('{}')), text)
    # Handle patterns like M_C, s_i, x_1, but not __bold__ or file_names
    text = re.sub(r'(?<![_a-z])([A-Za-z])_([A-Za-z0-9])(?![_a-z])', lambda m: m.group(1) + to_subscript(m.group(2)), text)
    # Handle multi-char subscript like s_{ij}
    text = re.sub(r'([A-Za-z])_\{([^}]+)\}', lambda m: m.group(1) + to_subscript(m.group(2)), text)
    # Arrow ->
    text = text.replace(' -> ', ' \u2192 ')
    # nabla
    text = text.replace('nabla(', '\u2207(')
    text = text.replace('nabla ', '\u2207')
    # delta
    text = re.sub(r'\bdelta([A-Z])', lambda m: '\u03B4' + m.group(1), text)
    # sigma
    text = re.sub(r'\bsigma\b', '\u03C3', text)
    # gamma
    text = re.sub(r'\bgamma\b', '\u03B3', text)
    # Dot product
    text = text.replace(' . ', ' \u00B7 ')
    return text


# ---------------------------------------------------------------------------
# Preprocessing: apply citations and math to full markdown text
# ---------------------------------------------------------------------------

def preprocess_markdown(md_text, bib):
    """Apply citation resolution and math notation to markdown."""
    md_text = resolve_citations(md_text, bib)
    md_text = convert_math_notation(md_text)
    return md_text


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------

def parse_markdown(text):
    """Parse markdown into structured blocks."""
    blocks = []
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if line.strip() == '---':
            blocks.append(('hr', ''))
            i += 1
            continue

        # Heading
        m = re.match(r'^(#{1,3})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            blocks.append(('heading', (level, m.group(2).strip())))
            i += 1
            continue

        # Blockquote (gather consecutive > lines)
        if line.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('>'):
                quote_lines.append(lines[i].lstrip('> ').rstrip())
                i += 1
            blocks.append(('quote', '\n'.join(quote_lines)))
            continue

        # Table
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            blocks.append(('table', table_lines))
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph (gather lines until empty or special)
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('>') and lines[i].strip() != '---':
            para_lines.append(lines[i].rstrip())
            i += 1
        if para_lines:
            blocks.append(('para', ' '.join(para_lines)))

    return blocks


def add_formatted_text(paragraph, text, base_size=11, base_font='Times New Roman', base_bold=False, base_italic=False):
    """Add text with inline markdown formatting (*italic*, **bold**)."""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(base_size)
            run.font.name = base_font
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(base_size)
            run.font.name = base_font
        else:
            run = paragraph.add_run(part)
            run.bold = base_bold
            run.italic = base_italic
            run.font.size = Pt(base_size)
            run.font.name = base_font


def add_page_number(section):
    """Add page number to footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'


def build_docx(blocks, md_text):
    """Build a print-ready .docx from parsed blocks."""
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # Page margins and page numbers
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.different_first_page_header_footer = True
        add_page_number(section)

    # --- Title page ---
    for _ in range(5):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('TRACTATUS FORMAE')
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacer

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('Ei substrat-uavhengig formteori')
    run.italic = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    for _ in range(3):
        doc.add_paragraph()

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_p.add_run('Iver Raknes Finne')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    inst_p = doc.add_paragraph()
    inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = inst_p.add_run('Institutt for design\nArkitektur- og designhøgskolen i Oslo')
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_page_break()

    # --- Body ---
    skip_title_block = True  # Skip the markdown title/author block
    first_heading_seen = False

    for btype, bdata in blocks:
        if btype == 'heading':
            level, text = bdata
            # Skip the first h1 (title) since we made a title page
            if level == 1 and skip_title_block and 'TRACTATUS' in text.upper():
                skip_title_block = False
                continue

            p = doc.add_paragraph()
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(12)
                add_formatted_text(p, text, base_size=18, base_bold=True)
            elif level == 2:
                if first_heading_seen:
                    # Page break before each main section (propositions)
                    if re.match(r'^\d\s', text):
                        doc.add_page_break()
                first_heading_seen = True
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
                add_formatted_text(p, text, base_size=14, base_bold=True)
            elif level == 3:
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                add_formatted_text(p, text, base_size=12, base_bold=True)

        elif btype == 'quote':
            # Skip if it's the subtitle/author block right after title
            if skip_title_block is False and not first_heading_seen:
                if 'substrat' in bdata.lower() or 'Iver' in bdata:
                    continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.right_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, bdata, base_size=10, base_italic=True)

        elif btype == 'para':
            # Skip subtitle/author lines from markdown
            if 'Ei substrat-uavhengig formteori' in bdata and not first_heading_seen:
                continue
            if 'Iver Raknes Finne' in bdata and not first_heading_seen:
                continue
            if bdata.strip().startswith('*Institutt for design') and not first_heading_seen:
                continue

            p = doc.add_paragraph()
            # Detect proposition lines
            prop_match = re.match(r'^(\*\*[\d.]+\*\*)\s+(\[.+?\])\s*(.*)', bdata)
            if prop_match:
                num = prop_match.group(1)[2:-2]
                tag = prop_match.group(2)
                rest = prop_match.group(3)
                p.paragraph_format.space_before = Pt(6)
                run = p.add_run(num + '  ')
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
                run = p.add_run(tag + '  ')
                run.italic = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                add_formatted_text(p, rest)
            else:
                add_formatted_text(p, bdata)

        elif btype == 'table':
            header = [c.strip() for c in bdata[0].split('|') if c.strip()]
            rows = []
            for row_line in bdata[2:]:
                cells = [c.strip() for c in row_line.split('|') if c.strip()]
                if cells:
                    rows.append(cells)
            if header and rows:
                table = doc.add_table(rows=1 + len(rows), cols=len(header))
                table.style = 'Table Grid'
                for j, h in enumerate(header):
                    cell = table.rows[0].cells[j]
                    cell.text = h
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                for i, row in enumerate(rows):
                    for j, val in enumerate(row):
                        if j < len(header):
                            cell = table.rows[i + 1].cells[j]
                            cell.text = val
                            for run in cell.paragraphs[0].runs:
                                run.font.size = Pt(9)
                                run.font.name = 'Times New Roman'

        elif btype == 'hr':
            # Thin separator line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT_DOCX)
    print(f"DOCX saved: {OUTPUT_DOCX}")


# ---------------------------------------------------------------------------
# PDF builder
# ---------------------------------------------------------------------------

def build_pdf(md_text):
    """Build a print-ready PDF via HTML/CSS with proper typography."""
    import markdown as md

    html_body = md.markdown(md_text, extensions=['tables', 'smarty'])

    # Post-process HTML: style proposition numbers
    html_body = re.sub(
        r'<strong>([\d.]+)</strong>\s*\[([^\]]+)\]',
        r'<span class="prop-num">\1</span> <span class="prop-tag">[\2]</span>',
        html_body
    )

    html_full = f"""<!DOCTYPE html>
<html lang="nn">
<head>
<meta charset="utf-8">
<style>
@page {{
    size: A4;
    margin: 2.5cm 2.5cm 3cm 2.5cm;
    @bottom-center {{
        content: counter(page);
        font-family: 'DejaVu Serif', Georgia, serif;
        font-size: 9pt;
        color: #333;
    }}
}}
@page :first {{
    @bottom-center {{ content: none; }}
}}
body {{
    font-family: 'DejaVu Serif', Georgia, serif;
    font-size: 11pt;
    line-height: 1.45;
    color: #000;
    hyphens: auto;
    -webkit-hyphens: auto;
}}
/* Title page */
.title-page {{
    page-break-after: always;
    text-align: center;
    padding-top: 8cm;
}}
.title-page h1 {{
    font-size: 28pt;
    letter-spacing: 0.05em;
    margin-bottom: 0.3em;
    text-align: center;
}}
.title-page .subtitle {{
    font-size: 14pt;
    font-style: italic;
    margin-bottom: 4cm;
}}
.title-page .author {{
    font-size: 14pt;
    margin-bottom: 0.3em;
}}
.title-page .institution {{
    font-size: 11pt;
    font-style: italic;
    color: #333;
}}
h1 {{
    font-size: 18pt;
    text-align: center;
    margin-top: 1.5em;
    margin-bottom: 0.8em;
    page-break-after: avoid;
}}
h2 {{
    font-size: 14pt;
    margin-top: 1.5em;
    margin-bottom: 0.6em;
    page-break-after: avoid;
    page-break-before: always;
}}
h2:first-of-type {{
    page-break-before: avoid;
}}
h3 {{
    font-size: 12pt;
    margin-top: 1em;
    margin-bottom: 0.4em;
    page-break-after: avoid;
}}
p {{
    margin: 0.35em 0;
    text-align: justify;
    orphans: 3;
    widows: 3;
}}
blockquote {{
    margin: 0.5em 0 0.5em 1.5cm;
    padding-left: 0.5cm;
    border-left: 2pt solid #999;
    font-size: 10pt;
    font-style: italic;
    color: #222;
}}
blockquote p {{
    text-align: left;
    margin: 0.2em 0;
}}
.prop-num {{
    font-weight: bold;
    font-size: 11pt;
}}
.prop-tag {{
    font-style: italic;
    font-size: 9pt;
    color: #555;
}}
table {{
    border-collapse: collapse;
    width: 100%;
    margin: 0.8em 0;
    font-size: 9pt;
    page-break-inside: avoid;
}}
th, td {{
    border: 1px solid #444;
    padding: 5px 8px;
    text-align: left;
}}
th {{
    background: #f0f0f0;
    font-weight: bold;
}}
hr {{
    border: none;
    border-top: 0.5pt solid #aaa;
    margin: 1.5em 0;
}}
code {{
    font-family: monospace;
    font-size: 10pt;
}}
</style>
</head>
<body>

<div class="title-page">
    <h1>TRACTATUS FORMAE</h1>
    <div class="subtitle">Ei substrat-uavhengig formteori</div>
    <div class="author">Iver Raknes Finne</div>
    <div class="institution">Institutt for design<br>Arkitektur- og designhøgskolen i Oslo</div>
</div>

{html_body}
</body>
</html>"""

    HTML(string=html_full).write_pdf(OUTPUT_PDF)
    print(f"PDF saved: {OUTPUT_PDF}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    # Load bibliography
    bib = parse_bibtex(BIB_FILE)
    print(f"Loaded {len(bib)} BibTeX entries")

    with open(INPUT, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Preprocess: resolve citations and math notation
    processed = preprocess_markdown(md_text, bib)

    # Build both formats
    blocks = parse_markdown(processed)
    build_docx(blocks, processed)
    build_pdf(processed)

    print("Done.")


if __name__ == '__main__':
    main()
