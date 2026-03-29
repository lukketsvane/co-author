#!/usr/bin/env python3
"""
Assemble transcribed Tractatus pages into a .docx file with proper formatting.
Reads raw transcription text files (batch_01.txt .. batch_10.txt) produced by
OCR agents, merges them in order, and builds a structured Word document that
preserves the hierarchical indentation of Wittgenstein's proposition numbering.
"""

import re
import sys
import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


BATCH_DIR = "/workspaces/NBNO.py/transcriptions"
OUTPUT = "/workspaces/NBNO.py/Tractatus_Logico-Philosophicus.docx"


def get_indent_level(prop_number: str) -> int:
    """Return indent level based on decimal structure of proposition number."""
    if '.' not in prop_number:
        return 0
    decimal_part = prop_number.split('.', 1)[1]
    return len(decimal_part)


def add_proposition(doc, number, text, italic_ranges=None):
    """Add a Tractatus proposition with proper indentation."""
    level = get_indent_level(number)
    indent_cm = level * 0.7  # 0.7cm per indent level

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm + 2.0)  # hanging indent for number
    p.paragraph_format.first_line_indent = Cm(-2.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    # Add the proposition number in bold
    num_run = p.add_run(number)
    num_run.bold = True
    num_run.font.size = Pt(11)
    num_run.font.name = 'Times New Roman'

    # Add tab separator
    p.add_run('\t')

    # Process text for italics (marked with *asterisks*)
    if text:
        parts = re.split(r'(\*[^*]+\*)', text)
        for part in parts:
            if part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    return p


def add_prose_paragraph(doc, text, indent=0, alignment=None, bold=False, italic=False, font_size=11):
    """Add a regular prose paragraph."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    if alignment == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Process text for italics
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        run.font.size = Pt(font_size)
        run.font.name = 'Times New Roman'

    return p


def add_footnote_paragraph(doc, text):
    """Add a footnote-style paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    # Process text for italics
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'

    return p


def add_parenthetical(doc, text, indent_level=0):
    """Add a parenthetical remark at the given indent level."""
    indent_cm = indent_level * 0.7 + 2.0
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    return p


def parse_and_build(doc, full_text):
    """Parse concatenated transcription text and build the document."""
    lines = full_text.split('\n')
    i = 0
    current_indent_level = 0  # track for parentheticals
    in_front_matter = True  # before main propositions start
    skip_page_numbers = True

    # Proposition pattern: optional tabs, then a number like 1, 1.1, 2.0121, etc., then tab, then text
    prop_pattern = re.compile(r'^(\t*)([\d]+(?:\.[\d]+)?)\t(.*)$')
    # Page marker
    page_pattern = re.compile(r'^=== PAGE (\d+) ===')
    # Footnote pattern
    footnote_pattern = re.compile(r'^\[FOOTNOTE (\d+): (.*)\]$')
    # Blank page
    blank_pattern = re.compile(r'^\[(BLANK PAGE|PAGE NUMBER ONLY)')
    # Section heading detection
    heading_pattern = re.compile(r'^[A-ZÆØÅ][A-ZÆØÅ\s\-]+$')

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip page markers
        if page_pattern.match(line):
            i += 1
            continue

        # Skip blank page markers
        if blank_pattern.match(line):
            i += 1
            continue

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Skip standalone page numbers (just a number at bottom of page)
        if re.match(r'^\d{1,3}$', line.strip()):
            i += 1
            continue

        # Check for proposition
        prop_match = prop_pattern.match(line)
        if prop_match:
            in_front_matter = False
            tabs = prop_match.group(1)
            number = prop_match.group(2)
            text = prop_match.group(3)

            # Gather continuation lines (lines that don't match a new proposition or marker)
            while i + 1 < len(lines):
                next_line = lines[i + 1].rstrip()
                if (not next_line.strip() or
                    page_pattern.match(next_line) or
                    prop_pattern.match(next_line) or
                    footnote_pattern.match(next_line) or
                    blank_pattern.match(next_line) or
                    next_line.strip().startswith('(') or
                    next_line.strip().startswith('[FOOTNOTE') or
                    re.match(r'^\d{1,3}$', next_line.strip()) or
                    re.match(r'^=== PAGE', next_line)):
                    break
                # This is a continuation line
                i += 1
                text += ' ' + next_line.strip()

            current_indent_level = get_indent_level(number)
            add_proposition(doc, number, text)
            i += 1
            continue

        # Check for footnote
        fn_match = footnote_pattern.match(line.strip())
        if fn_match:
            fn_text = f"{fn_match.group(1)} {fn_match.group(2)}"
            add_footnote_paragraph(doc, fn_text)
            i += 1
            continue

        # Check for parenthetical remark
        stripped = line.strip()
        if stripped.startswith('(') and not in_front_matter:
            # Gather full parenthetical
            paren_text = stripped
            while i + 1 < len(lines) and not paren_text.rstrip().endswith(')'):
                next_line = lines[i + 1].rstrip()
                if (page_pattern.match(next_line) or
                    prop_pattern.match(next_line) or
                    blank_pattern.match(next_line) or
                    re.match(r'^\d{1,3}$', next_line.strip())):
                    break
                i += 1
                paren_text += ' ' + next_line.strip()
            add_parenthetical(doc, paren_text, current_indent_level)
            i += 1
            continue

        # Regular prose text (front matter, afterword, etc.)
        # Check if it looks like a heading (all caps)
        if heading_pattern.match(stripped) and len(stripped) > 3:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # Regular paragraph
        add_prose_paragraph(doc, stripped)
        i += 1


def main():
    # Read all batch files in order
    batch_files = sorted([
        os.path.join(BATCH_DIR, f)
        for f in os.listdir(BATCH_DIR)
        if f.startswith('batch_') and f.endswith('.txt')
    ])

    if not batch_files:
        print("No batch files found in", BATCH_DIR)
        sys.exit(1)

    full_text = ""
    for bf in batch_files:
        with open(bf, 'r', encoding='utf-8') as f:
            full_text += f.read() + "\n"

    print(f"Read {len(batch_files)} batch files, total {len(full_text)} characters")

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title page
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('TRACTATUS\nLOGICO-PHILOSOPHICUS')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('Ludwig Wittgenstein')
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    trans_p = doc.add_paragraph()
    trans_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trans_p.paragraph_format.space_before = Pt(12)
    run = trans_p.add_run('Oversatt fra tysk og med etterord\nav Terje Ødegaard')
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    ded_p = doc.add_paragraph()
    ded_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ded_p.paragraph_format.space_before = Pt(24)
    run = ded_p.add_run('Til minne om min venn\nDavid H. Pinsent')
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Page break after title
    doc.add_page_break()

    # Parse and build the rest
    parse_and_build(doc, full_text)

    # Save
    doc.save(OUTPUT)
    print(f"Document saved to {OUTPUT}")


if __name__ == '__main__':
    main()
